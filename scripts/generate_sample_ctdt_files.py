"""
Tạo file mẫu CTĐT để smoke test extraction.
Chạy: python scripts/generate_sample_ctdt_files.py
Output: scripts/samples/
"""
from __future__ import annotations

import csv
import io
import os
import sys

_project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
samples_dir = os.path.join(_project_root, "scripts", "samples")
os.makedirs(samples_dir, exist_ok=True)


def generate_docx():
    """Mẫu 07 — Chương trình đào tạo (simplified)."""
    from docx import Document

    doc = Document()
    doc.add_heading("CHƯƠNG TRÌNH ĐÀO TẠO", level=1)
    doc.add_paragraph("Ngành: Công nghệ thông tin")
    doc.add_paragraph("Mã ngành: 7480201")
    doc.add_paragraph("Trình độ đào tạo: Đại học")
    doc.add_paragraph("Thời gian đào tạo: 4 năm")
    doc.add_paragraph("")

    doc.add_heading("I. MỤC TIÊU ĐÀO TẠO", level=2)
    doc.add_paragraph(
        "Đào tạo kỹ sư Công nghệ thông tin có phẩm chất chính trị, đạo đức, "
        "có kiến thức chuyên môn và kỹ năng thực hành nghề nghiệp."
    )

    doc.add_heading("II. CHUẨN ĐẦU RA", level=2)

    # Bảng CĐR
    cdr_data = [
        ["Mã CĐR", "Mô tả", "Mức độ (Bloom)"],
        ["C1", "Áp dụng kiến thức toán học, khoa học tự nhiên", "Áp dụng"],
        ["C2", "Phân tích, thiết kế hệ thống phần mềm", "Phân tích"],
        ["C3", "Lập trình, kiểm thử phần mềm", "Áp dụng"],
        ["C4", "Quản lý dự án CNTT", "Đánh giá"],
        ["C5", "Kỹ năng làm việc nhóm, giao tiếp", "Áp dụng"],
    ]
    t1 = doc.add_table(rows=len(cdr_data), cols=3)
    t1.style = "Table Grid"
    for i, row in enumerate(cdr_data):
        for j, val in enumerate(row):
            t1.rows[i].cells[j].text = val

    doc.add_heading("III. KHUNG CHƯƠNG TRÌNH ĐÀO TẠO", level=2)

    # Bảng học phần
    hp_data = [
        ["STT", "Mã HP", "Tên học phần", "Tổng TC", "LT", "TH", "BB/TC", "HP tiên quyết", "Học kỳ"],
        ["1", "CSE101", "Nhập môn CNTT", "3", "2", "1", "BB", "", "1"],
        ["2", "MAT101", "Toán cao cấp 1", "3", "3", "0", "BB", "", "1"],
        ["3", "CSE201", "Cấu trúc dữ liệu", "4", "3", "1", "BB", "CSE101", "2"],
        ["4", "CSE202", "Cơ sở dữ liệu", "3", "2", "1", "BB", "CSE101", "2"],
        ["5", "CSE301", "Lập trình web", "3", "2", "1", "TC", "CSE201", "3"],
        ["6", "CSE302", "Mạng máy tính", "3", "2", "1", "BB", "CSE101", "3"],
        ["7", "CSE401", "Đồ án tốt nghiệp", "10", "0", "10", "BB", "CSE301", "8"],
    ]
    t2 = doc.add_table(rows=len(hp_data), cols=9)
    t2.style = "Table Grid"
    for i, row in enumerate(hp_data):
        for j, val in enumerate(row):
            t2.rows[i].cells[j].text = val

    doc.add_paragraph("")
    doc.add_paragraph("Tổng số tín chỉ yêu cầu: 130 tín chỉ")

    doc.add_heading("IV. MA TRẬN MỤC TIÊU - CHUẨN ĐẦU RA", level=2)

    # Bảng ma trận
    matrix_data = [
        ["", "C1", "C2", "C3", "C4", "C5"],
        ["MT1", "3", "2", "", "", "1"],
        ["MT2", "", "3", "3", "2", ""],
        ["MT3", "", "", "", "3", "3"],
    ]
    t3 = doc.add_table(rows=len(matrix_data), cols=6)
    t3.style = "Table Grid"
    for i, row in enumerate(matrix_data):
        for j, val in enumerate(row):
            t3.rows[i].cells[j].text = val

    path = os.path.join(samples_dir, "Mau07_CTDT_CNTT.docx")
    doc.save(path)
    print(f"  ✅ {path} ({os.path.getsize(path):,} bytes)")


def generate_xlsx():
    """Danh mục học phần + ma trận CĐR dạng Excel."""
    import openpyxl

    wb = openpyxl.Workbook()

    # Sheet 1: Danh mục học phần
    ws1 = wb.active
    ws1.title = "DanhMucHocPhan"
    hp_rows = [
        ["STT", "Mã HP", "Tên học phần", "Tổng TC", "LT", "TH", "BB/TC", "HP tiên quyết", "Học kỳ"],
        [1, "CSE101", "Nhập môn CNTT", 3, 2, 1, "BB", None, 1],
        [2, "MAT101", "Toán cao cấp 1", 3, 3, 0, "BB", None, 1],
        [3, "CSE201", "Cấu trúc dữ liệu & Giải thuật", 4, 3, 1, "BB", "CSE101", 2],
        [4, "CSE202", "Cơ sở dữ liệu", 3, 2, 1, "BB", "CSE101", 2],
        [5, "CSE203", "Kiến trúc máy tính", 3, 2, 1, "BB", None, 2],
        [6, "CSE301", "Lập trình web", 3, 2, 1, "TC", "CSE201", 3],
        [7, "CSE302", "Mạng máy tính", 3, 2, 1, "BB", "CSE101", 3],
        [8, "CSE303", "Hệ điều hành", 3, 2, 1, "BB", "CSE203", 3],
        [9, "CSE401", "Trí tuệ nhân tạo", 3, 2, 1, "TC", "CSE201", 4],
        [10, "CSE402", "Học máy", 3, 2, 1, "TC", "CSE401", 5],
    ]
    for row in hp_rows:
        ws1.append(row)

    # Sheet 2: Ma trận CĐR - Học phần
    ws2 = wb.create_sheet("MaTran_CDR_HP")
    matrix = [
        ["CĐR \\ HP", "CSE101", "MAT101", "CSE201", "CSE202", "CSE301", "CSE302"],
        ["C1", 2, 3, 2, None, None, None],
        ["C2", None, None, 3, 3, 2, None],
        ["C3", 2, None, 3, 2, 3, None],
        ["C4", None, None, None, None, 2, 3],
        ["C5", 1, None, 2, None, 2, 2],
    ]
    for row in matrix:
        ws2.append(row)

    # Sheet 3: Empty (should be skipped)
    wb.create_sheet("TrangTrong")

    path = os.path.join(samples_dir, "DanhMuc_HocPhan_CNTT.xlsx")
    wb.save(path)
    print(f"  ✅ {path} ({os.path.getsize(path):,} bytes)")


def generate_csv():
    """Khảo sát sinh viên dạng CSV."""
    rows = [
        ["STT", "MSSV", "Họ tên", "Đánh giá chung", "Mức hài lòng", "Góp ý"],
        ["1", "20T1080001", "Nguyễn Văn A", "Tốt", "4", "Chương trình phù hợp"],
        ["2", "20T1080002", "Trần Thị B", "Khá", "3", "Cần thêm môn thực hành"],
        ["3", "20T1080003", "Lê Văn C", "Tốt", "5", ""],
        ["4", "20T1080004", "Phạm Thị D", "Trung bình", "3", "Thiếu môn tự chọn"],
        ["5", "20T1080005", "Hoàng Văn E", "Tốt", "4", "Nên cập nhật giáo trình"],
    ]
    path = os.path.join(samples_dir, "KhaoSat_SV_CNTT.csv")
    with open(path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerows(rows)
    print(f"  ✅ {path} ({os.path.getsize(path):,} bytes)")


if __name__ == "__main__":
    print("Generating sample CTĐT files...")
    generate_docx()
    generate_xlsx()
    generate_csv()
    print()
    print("Done! Chạy smoke test:")
    print(f"  python scripts/smoke_extract_ctdt.py {os.path.join(samples_dir, 'Mau07_CTDT_CNTT.docx')}")
    print(f"  python scripts/smoke_extract_ctdt.py {os.path.join(samples_dir, 'DanhMuc_HocPhan_CNTT.xlsx')}")
    print(f"  python scripts/smoke_extract_ctdt.py {os.path.join(samples_dir, 'KhaoSat_SV_CNTT.csv')}")
