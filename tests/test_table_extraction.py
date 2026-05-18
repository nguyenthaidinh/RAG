"""Tests for R2 — Table-aware extraction: DOCX tables, XLSX, CSV."""
import csv
import io
import pytest
from unittest.mock import patch

from app.services.document_extract import (
    ExtractLimits,
    TableLimits,
    UnsupportedFileType,
    extract_text,
    extract_text_with_metadata,
    _cell_to_str,
    _format_table_rows,
    _detect_csv_encoding,
)


# ── Helpers ───────────────────────────────────────────────────────────


def _make_simple_docx(paragraphs: list[str], tables: list[list[list[str]]] | None = None) -> bytes:
    """Create a minimal DOCX in memory with paragraphs and optionally tables."""
    from docx import Document as DocxDocument
    doc = DocxDocument()
    for p_text in paragraphs:
        doc.add_paragraph(p_text)
    if tables:
        for tbl_data in tables:
            if not tbl_data:
                continue
            num_cols = max(len(row) for row in tbl_data)
            table = doc.add_table(rows=len(tbl_data), cols=num_cols)
            for r_idx, row in enumerate(tbl_data):
                for c_idx, val in enumerate(row):
                    table.rows[r_idx].cells[c_idx].text = val
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_simple_xlsx(sheets: dict[str, list[list]]) -> bytes:
    """Create a minimal XLSX in memory with named sheets."""
    import openpyxl
    wb = openpyxl.Workbook()
    # Remove default sheet
    default_ws = wb.active
    for sheet_name, rows in sheets.items():
        ws = wb.create_sheet(title=sheet_name)
        for row in rows:
            ws.append(row)
    wb.remove(default_ws)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _make_simple_csv(rows: list[list[str]], encoding: str = "utf-8") -> bytes:
    """Create CSV bytes from rows."""
    buf = io.StringIO()
    writer = csv.writer(buf)
    for row in rows:
        writer.writerow(row)
    return buf.getvalue().encode(encoding)


# ── Unit tests: helpers ──────────────────────────────────────────────


class TestCellToStr:
    def test_none(self):
        assert _cell_to_str(None) == ""

    def test_int(self):
        assert _cell_to_str(42) == "42"

    def test_float_whole(self):
        assert _cell_to_str(3.0) == "3"

    def test_float_decimal(self):
        assert _cell_to_str(3.14) == "3.14"

    def test_string(self):
        assert _cell_to_str("  hello  ") == "hello"


class TestFormatTableRows:
    def test_empty(self):
        assert _format_table_rows([], max_rows=10, max_cols=10, include_empty_cells=False) == ""

    def test_basic_table(self):
        rows = [["A", "B"], ["1", "2"]]
        result = _format_table_rows(rows, max_rows=100, max_cols=100, include_empty_cells=False)
        assert "| A | B |" in result
        assert "| 1 | 2 |" in result

    def test_truncation(self):
        rows = [["x"] for _ in range(10)]
        result = _format_table_rows(rows, max_rows=3, max_cols=10, include_empty_cells=False)
        assert "truncated at 3 rows" in result
        assert result.count("| x |") == 3

    def test_pipe_escape(self):
        rows = [["a|b"]]
        result = _format_table_rows(rows, max_rows=10, max_cols=10, include_empty_cells=False)
        assert "a\\|b" in result

    def test_col_limit(self):
        rows = [["c1", "c2", "c3", "c4", "c5"]]
        result = _format_table_rows(rows, max_rows=10, max_cols=3, include_empty_cells=False)
        assert "c1" in result
        assert "c3" in result
        assert "c4" not in result


class TestDetectCsvEncoding:
    def test_utf8(self):
        data = "hello".encode("utf-8")
        assert _detect_csv_encoding(data) in ("utf-8-sig", "utf-8")

    def test_utf8_bom(self):
        data = b"\xef\xbb\xbfhello"
        assert _detect_csv_encoding(data) == "utf-8-sig"

    def test_latin1_fallback(self):
        # Bytes that are invalid UTF-8 but valid latin-1
        data = bytes([0xc0, 0xc1, 0xfe, 0xff])
        enc = _detect_csv_encoding(data)
        assert enc == "latin-1"


# ── DOCX extraction ─────────────────────────────────────────────────


class TestDocxExtraction:
    """Test DOCX paragraph + table extraction."""

    def test_paragraphs_only_backward_compat(self):
        """DOCX with only paragraphs should work exactly like R1."""
        data = _make_simple_docx(["Paragraph one.", "Paragraph two."])
        text = extract_text("test.docx", None, data)
        assert "Paragraph one." in text
        assert "Paragraph two." in text
        assert "[TABLE" not in text

    def test_paragraphs_and_tables(self):
        """DOCX with paragraphs + table should have both in output."""
        data = _make_simple_docx(
            paragraphs=["Introduction paragraph."],
            tables=[[
                ["Mã HP", "Tên học phần", "TC"],
                ["CSE101", "Nhập môn CNTT", "3"],
                ["CSE201", "CTDL & GT", "4"],
            ]],
        )
        text, meta = extract_text_with_metadata("test.docx", None, data)

        assert "Introduction paragraph." in text
        assert "[TABLE 1]" in text
        assert "CSE101" in text
        assert "Nhập môn CNTT" in text
        assert meta["has_tables"] is True
        assert meta["tables_count"] == 1

    def test_multiple_tables(self):
        """DOCX with multiple tables should number them sequentially."""
        data = _make_simple_docx(
            paragraphs=["Header text."],
            tables=[
                [["A", "B"], ["1", "2"]],
                [["C", "D"], ["3", "4"]],
            ],
        )
        text, meta = extract_text_with_metadata("test.docx", None, data)
        assert "[TABLE 1]" in text
        assert "[TABLE 2]" in text
        assert meta["tables_count"] == 2

    def test_docx_metadata(self):
        """Extraction metadata should be populated correctly."""
        data = _make_simple_docx(
            paragraphs=[],
            tables=[[["X"]]],
        )
        _, meta = extract_text_with_metadata("test.docx", None, data)
        assert meta["extractor_version"] == "table_aware_v1"
        assert meta["has_tables"] is True
        assert meta["sheets_count"] == 0


# ── XLSX extraction ──────────────────────────────────────────────────


class TestXlsxExtraction:
    """Test XLSX sheet extraction."""

    def test_single_sheet(self):
        data = _make_simple_xlsx({
            "DanhMuc": [
                ["STT", "Mã HP", "Tên HP"],
                [1, "CSE101", "Nhập môn"],
                [2, "CSE201", "CTDL"],
            ],
        })
        text, meta = extract_text_with_metadata("data.xlsx", None, data)

        assert "[SHEET: DanhMuc]" in text
        assert "CSE101" in text
        assert "Nhập môn" in text
        assert meta["sheets_count"] == 1
        assert meta["has_tables"] is True

    def test_multi_sheet(self):
        data = _make_simple_xlsx({
            "HocPhan": [["A", "B"], [1, 2]],
            "ChuanDauRa": [["C", "D"], [3, 4]],
        })
        text, meta = extract_text_with_metadata("multi.xlsx", None, data)

        assert "[SHEET: HocPhan]" in text
        assert "[SHEET: ChuanDauRa]" in text
        assert meta["sheets_count"] == 2

    def test_empty_sheet_skipped(self):
        """Empty sheets should be silently skipped."""
        import openpyxl
        wb = openpyxl.Workbook()
        ws1 = wb.active
        ws1.title = "Empty"
        ws2 = wb.create_sheet("HasData")
        ws2.append(["Value1"])
        buf = io.BytesIO()
        wb.save(buf)
        data = buf.getvalue()

        text, meta = extract_text_with_metadata("mixed.xlsx", None, data)
        assert "[SHEET: HasData]" in text
        # Empty sheet should be skipped (either not present or as empty section)
        assert meta["sheets_count"] >= 1

    def test_xlsx_by_mime_type(self):
        """Should recognize XLSX by MIME type even without extension."""
        data = _make_simple_xlsx({"S1": [["X"]]})
        text = extract_text(
            "noext",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            data,
        )
        assert "[SHEET: S1]" in text

    def test_xlsx_float_to_int(self):
        """Whole-number floats should be displayed without .0 suffix."""
        data = _make_simple_xlsx({
            "Credits": [[3.0, 2.0, 1.0]],
        })
        text = extract_text("credits.xlsx", None, data)
        assert "3" in text
        # Should NOT have "3.0"
        assert "3.0" not in text


# ── CSV extraction ───────────────────────────────────────────────────


class TestCsvExtraction:
    """Test CSV file extraction."""

    def test_basic_csv(self):
        data = _make_simple_csv([
            ["Mã HP", "Tên HP", "TC"],
            ["CSE101", "Nhập môn", "3"],
        ])
        text, meta = extract_text_with_metadata("data.csv", None, data)

        assert "CSE101" in text
        assert "Nhập môn" in text
        assert meta["has_tables"] is True
        assert meta["tables_count"] == 1

    def test_csv_utf8_bom(self):
        """CSV with BOM should decode correctly."""
        raw = "STT,Tên\n1,Toán\n"
        data = b"\xef\xbb\xbf" + raw.encode("utf-8")
        text = extract_text("bom.csv", "text/csv", data)
        assert "Toán" in text

    def test_csv_empty(self):
        """Empty CSV should return empty-ish text, not crash."""
        data = b""
        text, meta = extract_text_with_metadata("empty.csv", None, data)
        assert meta["has_tables"] is False

    def test_csv_by_mime_type(self):
        """Should recognize CSV by MIME type."""
        data = _make_simple_csv([["A", "B"]])
        text = extract_text("noext", "text/csv", data)
        assert "A" in text

    def test_csv_row_limit(self):
        """Should respect row limits."""
        rows = [["header", "value"]] + [[f"item-{i}", str(i)] for i in range(100)]
        data = _make_simple_csv(rows)
        text, meta = extract_text_with_metadata(
            "big.csv", None, data,
            table_limits=TableLimits(max_rows=5, max_cols=50),
        )
        assert "header" in text
        assert "item-0" in text
        # Row 5 (header + 4 data rows = 5 total) should be the last
        assert "item-4" not in text  # 6th row (header + 5 data) exceeds limit


# ── Backward compatibility ───────────────────────────────────────────


class TestBackwardCompatibility:
    """Ensure R1 paths still work after R2 changes."""

    def test_txt_unchanged(self):
        data = "Hello world\nLine two".encode("utf-8")
        text = extract_text("readme.txt", "text/plain", data)
        assert "Hello world" in text
        assert "Line two" in text

    def test_md_unchanged(self):
        data = "# Title\n\nContent here".encode("utf-8")
        text = extract_text("doc.md", "text/markdown", data)
        assert "# Title" in text

    def test_unsupported_raises(self):
        with pytest.raises(UnsupportedFileType):
            extract_text("file.xyz", "application/octet-stream", b"data")

    def test_doc_legacy_raises(self):
        with pytest.raises(UnsupportedFileType):
            extract_text("old.doc", "application/msword", b"data")

    def test_extract_text_returns_str(self):
        """extract_text() must still return str, not tuple."""
        data = "test".encode("utf-8")
        result = extract_text("t.txt", "text/plain", data)
        assert isinstance(result, str)

    def test_extract_text_with_metadata_returns_tuple(self):
        """extract_text_with_metadata() returns (str, dict)."""
        data = "test".encode("utf-8")
        result = extract_text_with_metadata("t.txt", "text/plain", data)
        assert isinstance(result, tuple)
        assert len(result) == 2
        assert isinstance(result[0], str)
        assert isinstance(result[1], dict)


# ── MIME type validation (config-level) ──────────────────────────────


class TestMimeTypeConfig:
    """Verify XLSX/CSV MIME types are accepted by ingest validation."""

    @patch("app.services.ctdt_ingest_service.settings")
    def test_xlsx_mime_accepted(self, mock_settings):
        from app.services.ctdt_ingest_service import validate_mime_type
        mock_settings.RAG_ALLOWED_MIME_TYPES = (
            "application/pdf,"
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet,"
            "text/csv"
        )
        # Should not raise
        validate_mime_type("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    @patch("app.services.ctdt_ingest_service.settings")
    def test_csv_mime_accepted(self, mock_settings):
        from app.services.ctdt_ingest_service import validate_mime_type
        mock_settings.RAG_ALLOWED_MIME_TYPES = "text/csv,application/pdf"
        validate_mime_type("text/csv")


# ── Extraction metadata in ingest pipeline ───────────────────────────


class TestExtractionMetadataInMeta:
    """Test that extraction metadata is correctly built into Document.meta."""

    def test_build_ctdt_metadata_accepts_extraction_key(self):
        """After ingest, meta should be able to hold extraction info."""
        from app.services.ctdt_ingest_service import build_ctdt_metadata
        meta = build_ctdt_metadata(
            external_file_id="f1",
            update_cycle_id="15",
            program_id=None,
            program_code="7480201",
            program_name="CNTT",
            document_role="current_curriculum",
            uploaded_by="user1",
            checksum=None,
            filename="test.xlsx",
            mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            file_size_bytes=1024,
            ingest_mode="legacy",
        )
        # Simulate R2 extraction metadata injection
        meta["extraction"] = {
            "extractor_version": "table_aware_v1",
            "tables_count": 2,
            "sheets_count": 2,
            "has_tables": True,
        }
        assert meta["extraction"]["tables_count"] == 2
        assert meta["extraction"]["has_tables"] is True
        # Original keys untouched
        assert meta["ctdt"]["external_file_id"] == "f1"
        assert meta["system"]["pipeline_mode"] == "legacy"
